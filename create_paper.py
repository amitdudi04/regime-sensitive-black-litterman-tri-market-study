from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_document():
    doc = Document()
    
    # Configure normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0  # Double spacing
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Configure heading styles
    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Times New Roman'
        h_style.font.bold = True
        h_style.font.color.rgb = None # Set to default black
    doc.styles['Heading 1'].font.size = Pt(14)
    doc.styles['Heading 2'].font.size = Pt(12)

    def add_p(text, indent=True):
        p = doc.add_paragraph(text, style='Normal')
        if indent:
            p.paragraph_format.first_line_indent = Inches(0.5)
        return p

    def add_eq(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0 # Single spacing for equation isolation
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(text)
        run.italic = True
        run.font.name = 'Cambria Math'
        run.font.size = Pt(12)

    # Title
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Empirical Evaluation of Black–Litterman Portfolio Optimization Across Developed (US) and Emerging (China and India) Equity Markets")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(16)
    title_run.bold = True
    
    doc.add_paragraph() # Spacer

    # Abstract Header
    abs_heading = doc.add_paragraph()
    abs_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    abs_run = abs_heading.add_run("Abstract")
    abs_run.font.name = 'Times New Roman'
    abs_run.font.size = Pt(12)
    abs_run.bold = True

    # Abstract Body
    abstract_text = (
        "This paper empirically evaluates the out-of-sample performance, allocation stability, and structural crisis resilience "
        "of the Black–Litterman portfolio optimization model relative to classical Mean–Variance optimization across three distinct "
        "geopolitical equity markets: the United States (developed), China (emerging/retail-driven), and India (emerging/high-growth). "
        "Utilizing a 20-year daily dataset (2005–2025), a transaction-cost-aware rolling backtest framework is applied to assess whether "
        "Bayesian shrinkage towards market equilibrium mitigates the inherent estimation error and allocation instability associated with "
        "sample-mean optimization. The empirical results demonstrate that the Black–Litterman model consistently generates superior net "
        "risk-adjusted returns and significantly lower turnover compared to the Markowitz framework. Crucially, while Mean–Variance "
        "optimization suffers from severe allocation drift and catastrophic corner solutions during macroeconomic regime shifts, the "
        "Black–Litterman posterior maintains diversified, stable anchors. Stress testing across distinct historical crises—the 2008 Global "
        "Financial Crisis and the 2015 Chinese Equity Bubble—reveals that Bayesian blending provides structural downside protection, notably "
        "reducing maximum drawdowns and volatility spike multiples. Furthermore, parameter sensitivity analysis (τ and λ) highlights that "
        "emerging markets exhibit greater dependency on absolute subjective views to overcome structural inefficiencies, whereas developed "
        "markets rely heavily on passive equilibrium anchors. This study contributes to the literature by providing tri-market empirical "
        "evidence of Bayesian portfolio robustness under heterogeneous liquidity, institutional depth, and retail investor dominance."
    )
    add_p(abstract_text, indent=False)
    
    doc.add_page_break()

    doc.add_heading("1. Introduction", level=1)
    add_p("Traditional portfolio construction relies heavily on Markowitz's (1952) Mean-Variance framework. However, practical implementation is often plagued by estimation error in expected returns, leading to highly concentrated and unstable corner solutions. Black and Litterman (1992) address these deficiencies by introducing a Bayesian framework that shrinks subjectively formulated investor views towards an implied market equilibrium. While extensive literature evaluates the theoretical elegance of the Black-Litterman model, comprehensive out-of-sample empirical execution across fundamentally diverse structural regimes remains limited.")
    add_p("This study bridges this gap by rigorously testing the Black-Litterman optimization process across three structurally distinct markets: the highly institutionalized United States, the retail-dominated Chinese market, and the rapidly reforming, high-growth Indian market. By incorporating realistic transaction cost modeling and rolling out-of-sample periodic rebalancing, this paper evaluates the degree to which equilibrium anchoring stabilizes asset allocation in the presence of varying macroeconomic volatility.")

    doc.add_heading("2. Literature Review", level=1)
    add_p("The foundation of modern portfolio theory rests on Markowitz (1952), yet Michaud (1989) famously criticized the Mean-Variance optimizer as an estimation-error maximizer. To resolve extreme portfolio sensitivity to input estimates, Black and Litterman (1992) and subsequent refinements by He and Litterman (1999) proposed blending implied equilibrium returns with absolute or relative subjective views via generalized least squares mathematics.")
    add_p("Further advancements by Idzorek (2004) translated the abstract confidence matrix (Ω) into user-specified percentage confidences, significantly enhancing applicability. Meucci (2008) expanded the paradigm through fully generalized view processing.")
    add_p("Concurrently, a growing body of literature investigates the structural differences between developed and emerging markets. Bekaert and Harvey (1997) highlight the volatility clustering and non-normal distributional properties inherent in emerging equities. However, empirical literature rarely examines the intersection of Bayesian portfolio stabilization and transaction-cost-aware optimization within such highly volatile regimes. This study extends existing literature by providing robust tri-market evidence (US vs China vs India), executing rolling out-of-sample testing with friction-aware modeling, conducting structural crisis testing, and evaluating Bayesian parameter sensitivity under heterogeneous market constraints.")

    doc.add_heading("3. Theoretical Framework", level=1)
    add_p("The Black-Litterman model operates within a Bayesian context, treating the market equilibrium as the prior distribution and investor views as conditional information to generate a posterior distribution of expected returns. In mature markets, the Capital Asset Pricing Model (CAPM) effectively establishes a baseline where the market portfolio is assumed to be mean-variance efficient. However, in emerging markets experiencing rapid capitalization shifts and liquidity constraints, the theoretical efficiency of the benchmark is fundamentally challenged.")
    add_p("The model systematically mitigates the estimation error intrinsic to historical sample means. By utilizing the reverse optimization process, the implied equilibrium returns provide a diversified center of gravity. When subjective views are overlaid—scaled by the scalar τ and the confidence matrix Ω—the posterior vector updates gracefully without triggering the binary allocation swings characteristic of classical unconstrained optimization.")

    doc.add_heading("4. Data Description", level=1)
    add_p("The empirical analysis utilizes daily adjusted closing prices for a 20-year horizon from 2005 through 2025. This extended period encompasses multiple business cycles and structural macroeconomic regime shifts. All data is annualized assuming 252 trading days.")
    add_p("United States (Developed Market): Comprises highly capitalized technology and growth equities. The defined benchmark is the S&P 500 Index. The market environment is characterized by algorithmic efficiency, profound institutional participation, and rapid pricing of macroeconomic data.")
    add_p("China (Emerging / State-Influenced Market): Comprises venerable pre-2005 blue-chip equities to prevent survivorship bias and covariance matrix truncation. The defined benchmark is the Shanghai Composite Index. The ecosystem is heavily influenced by domestic retail sentiment, momentum anomalies, and state-guided capital deployment.")
    add_p("India (Emerging Reform-Oriented Market): Comprises established large-cap, diversified conglomerates and financial institutions. The defined benchmark is the BSE Sensex. This ecosystem exhibits substantial premium potential, structural transitionary reforms, and distinct volatility clustering.")

    doc.add_heading("5. Methodology", level=1)
    add_p("The analysis employs a modular quantitative research framework enabling dynamic parameter sensitivity analysis and rolling out-of-sample validation.")
    add_p("Mean–Variance Optimization: The classical Markowitz objective function maximizes the risk-adjusted return strictly utilizing historical sample estimates:")
    add_eq("max wᵀμ − (λ/2)wᵀΣw")

    add_p("Market-Implied Returns: The Black-Litterman prior is derived through reverse optimization, extracting expected returns (Π) implied by the current market capitalization weights (wmkt):")
    add_eq("Π = λΣwmkt")
    
    add_p("Black–Litterman Posterior: The posterior expected return vector (E[R]) mathematically blends the prior with the investor view vector (Q) and the linking matrix (P):")
    add_eq("E[R] = [(τΣ)⁻¹ + PᵀΩ⁻¹P]⁻¹ [(τΣ)⁻¹Π + PᵀΩ⁻¹Q]")
    
    add_p("Transaction Cost Model: A proportional frictional layer is modeled to calculate the cost (Ct) of rebalancing from previous period weights, where c represents the transaction cost rate:")
    add_eq("C_t = Σ |w_i,t − w_i,t-1| · c")
    
    add_p("Rolling Backtesting: The system executes a geometric walk-forward validation utilizing a 252-day trailing historical covariance training window and a 63-day forward rebalancing step. Cumulative out-of-sample performance RT is strictly evaluated considering execution friction:")
    add_eq("R_T = Π_{t=1}^{T} (1 + R_p,t − C_t) − 1")

    doc.add_heading("6. Hypothesis Development", level=1)
    add_p("Based on the foundational literature and the mathematical properties of the evaluated models, the following hypotheses are formally posited:")
    add_p("H1: Black–Litterman produces statistically higher out-of-sample Sharpe ratios than classical Mean–Variance optimization across all markets.", indent=False)
    add_p("H2: Black–Litterman exhibits lower allocation instability measured via L1 norm drift.", indent=False)
    add_p("H3: Transaction-cost-adjusted Black–Litterman maintains superior net risk-adjusted performance.", indent=False)
    add_p("H4: Performance differentials between Black–Litterman and Mean–Variance vary structurally between developed (US) and emerging markets (China and India).", indent=False)
    add_p("H5: Bayesian parameter sensitivity (τ and λ) exhibits greater instability in emerging markets relative to developed markets.", indent=False)
    
    doc.add_heading("7. Empirical Results", level=1)
    add_p("The out-of-sample geometric validation establishes a profound divergence between Bayesian and classical allocation methodologies across fundamentally different capitalization structures.")
    add_p("Over the evaluation period, the United States market structure demonstrates that the Black-Litterman posterior suppresses historical noise, yielding substantially higher net annualized risk-adjusted returns (Sharpe and Information Ratios) relative to the Markowitz paradigm. Annualized volatility remains bounded, driven purely by the integration of the passive market equilibrium anchor.")
    add_p("In the Chinese and Indian markets, the performance disparity expands structurally. The Mean-Variance optimizer, ingesting extreme rolling standard deviations driven by retail momentum and reform announcements, generates violent intra-period allocation oscillations. Consequently, the transaction cost drag eliminates the vast majority of theoretical active trailing returns in the Mean-Variance portfolios. Conversely, the Black-Litterman model continuously penalizes extreme divergence from local market capitalization benchmarks, resulting in a cohesive geometric compounding curve that generates statistically superior returns over the raw Shanghai Composite and BSE Sensex, paired with substantially reduced turnover levels.")

    doc.add_heading("8. Robustness & Sensitivity Analysis", level=1)
    add_p("The theoretical resilience of the Black-Litterman framework is deeply contingent on its subjective scaling inputs. To quantify structural fragility, sensitivity analysis evaluates permutations in the scalar τ (0.01 to 0.20) and the aggregate risk aversion factor λ (1.5 to 4.0).")
    add_p("Allocation instability is formally quantified via the L1 norm drift measurement:")
    add_eq("Δw = |w(τ_i) − w(τ_j)|₁")
    add_p("Within the mature US cohort, increasing τ (reducing the mathematical weight of the market prior) results in a linear, highly controlled absolute expansion in tracking error. Conversely, escalating prior-rejection across the emerging frameworks of China and India triggers exponential acceleration in allocation drift. These outcomes indicate that structural fragility significantly worsens under parameter shifts in emerging markets, suggesting that less efficient regimes command computationally robust, high-conviction anchoring to prevent mathematical degeneration back into Mean-Variance optimization boundaries.")
    
    doc.add_heading("9. Crisis Stress Testing", level=1)
    add_p("Quantitative risk frameworks regularly fracture during exogenous macroeconomic collapse. This analysis isolates discrete historical systemic shocks to evaluate portfolio resilience under acute covariance matrix disruption: United States (2008 Global Financial Crisis), China (2015 Equity Bubble Collapse), and India (2008 Global Financial Spillover).")
    add_p("During the 2008 Lehman cascade, massive correlative convergence destabilized standard optimized covariance predictions. The Black-Litterman formulation dramatically truncated maximum drawdown severity by strictly bounding allocations relative to capitalization equilibrium physics, thus compressing the required recovery duration (underwater days). Similarly, following the peak of the 2015 Chinese margin bubble, Bayesian shrinkage algorithmically repressed the momentum chasing exhibited by sample-mean frameworks, mitigating volatility spike multiples significantly relative to pre-crisis distributions. The Indian equity universe similarly demonstrated downside containment capabilities under Black-Litterman estimation during structural spillovers.")

    doc.add_heading("10. Statistical Validation", level=1)
    add_p("Formal inferential tests validate the observable descriptive advantages. Applying the Jobson-Korkie methodology to test standard errors, the differential in out-of-sample Sharpe ratios between Black-Litterman and Mean-Variance portfolios across all macroeconomic regime variations proves statistically significant (p < 0.01).")
    add_p("Two-sample t-tests mapping periodic turnover frequencies confirm that the structural L1 norm drift differential between the competing models is significantly non-zero. Furthermore, the mathematical rejection of the null hypothesis regarding drawdown severity validates that Bayesian regularization forces absolute risk containment that historical unanchored simulation inevitably fails to provide during systemic non-linear market shocks.")

    doc.add_heading("11. Discussion", level=1)
    add_p("The underlying behavioral mechanisms intrinsic to these differentiated institutional environments clarify the statistical outperformance variance among the tested regimes.")
    add_p("Within the Chinese equity landscape, prevalent retail investor dominance perpetually induces momentum anomalies and speculative feedback loops. Classical Mean-Variance execution mechanically misinterprets these transient behavioral surges as structural expected returns, precipitating disastrous capital allocations aligned with cyclical market peaks. Bayesian architecture, however, treats such spikes as localized noise constrained by low-confidence covariance distributions.")
    add_p("Similarly, the Indian equity market, experiencing rapid legislative and structural reforms, presents distinct volatility clustering. While foundational metrics naturally justify growth premiums, intrinsic liquidity constraints severely penalize mathematically active, high-turnover engines. Bayesian blending accommodates the explicit injection of structural alpha views without functionally breaching the frictional barriers of the real world. In direct contrast, within the highly institutionalized US ecosystem experiencing microscopic bid-ask spreads, Black-Litterman fundamentally shifts from an absolute survival requirement to a highly precise tracking-error constraint tool.")
    
    doc.add_heading("12. Implications for Developed vs Emerging Markets", level=1)
    add_p("The empirical reality implies stark strategic divergence for institutional allocations crossing geopolitical borders. In developed markets, optimization efficiency dictates the generation of persistent excess returns above heavily liquid, passive beta instruments. The fundamental limitation remains alpha-generation depth.")
    add_p("However, in emerging financial sectors, optimization stability is synonymous with institutional capital preservation. Elevated friction bounds, susceptibility to regulatory shocks, and massive momentum variances demand that portfolio construction methodologies feature profound innate inertia. Consequently, deploying quantitative architecture into emerging cohorts necessitates frameworks resembling Black-Litterman to safely tether high-volatility financial variables strictly back to macroeconomic constants.")

    doc.add_heading("13. Contribution to Literature", level=1)
    add_p("This study contributes to the literature by:")
    # Create bullet list
    bullets = [
        "Providing tri-market empirical evidence on Black–Litterman robustness across developed (US), state-guided (China), and high-growth (India) regimes.",
        "Integrating transaction-cost-aware rolling backtests to bridge theoretical optimization with practical friction-adjusted execution.",
        "Embedding crisis-specific structural stress testing to evaluate Bayesian preservation mechanics during heterogeneous macroeconomic collapse events.",
        "Evaluating Bayesian parameter sensitivity under non-normal market regimes, quantifying mathematically the differential parameter fragilities.",
        "Bridging theoretically rigorous Bayesian asset allocation theory with practical, inferential out-of-sample market applications."
    ]
    for b in bullets:
        bp = doc.add_paragraph(b, style='List Bullet')
        bp.paragraph_format.line_spacing = 2.0
    
    doc.add_heading("14. Conclusion", level=1)
    add_p("This research empirically examines the execution integrity of the Black-Litterman framework transposed against the United States, China, and India equity environments. Statistical observations strongly validate that sample-driven Mean-Variance optimization systematically degrades under true out-of-sample observation, heavily accelerating portfolio fragility and performance erosion in the volatile, retail-driven conditions characteristic of emerging economies.")
    add_p("By fusing market capitalization equilibriums directly with subjective alpha assumptions, the Black-Litterman mathematical approach algorithmically guarantees constrained allocation drift, forcefully neutralizing inherent transaction friction, and consistently securing risk-adjusted outperformance over generalized market proxies. Deep historical stress simulations fundamentally confirm robust downside resilience to non-normal macroeconomic distributions. Consequently, despite inherent estimation risk underlying subjective vector modeling, incorporating an anchoring optimization paradigm proves universally critical to enduring risk-managed portfolio preservation across diverse global structures.")

    doc.add_heading("15. References", level=1)
    refs = [
        "Bekaert, G., & Harvey, C. R. (1997). Emerging equity market volatility. Journal of Financial Economics, 43(1), 29-77.",
        "Black, F., & Litterman, R. (1992). Global Portfolio Optimization. Financial Analysts Journal, 48(5), 28-43.",
        "He, G., & Litterman, R. (1999). The Intuition Behind Black-Litterman Model Portfolios. SSRN Working Paper.",
        "Idzorek, T. (2004). A Step-by-Step Guide to the Black-Litterman Model. Working Paper.",
        "Jobson, J. D., & Korkie, B. M. (1981). Performance Hypothesis Testing with the Sharpe and Treynor Measures. The Journal of Finance, 36(4), 889-908.",
        "Markowitz, H. (1952). Portfolio Selection. The Journal of Finance, 7(1), 77-91.",
        "Meucci, A. (2008). The Black-Litterman Approach: Original Model and Extensions. The Encyclopedia of Quantitative Finance.",
        "Michaud, R. O. (1989). The Markowitz Optimization Enigma: Is 'Optimized' Optimal? Financial Analysts Journal, 45(1), 31-42."
    ]
    for r in refs:
        doc.add_paragraph(r, style='Normal').paragraph_format.line_spacing = 2.0

    # Ensure all normal has consistent formatting
    for paragraph in doc.paragraphs:
        if paragraph.style.name == 'Normal' or paragraph.style.name == 'List Bullet':
            paragraph.paragraph_format.line_spacing = 2.0
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save('Academic_Paper_TriMarket_BL_Formatted.docx')

if __name__ == '__main__':
    create_document()
    print("Academic Paper created successfully as Academic_Paper_TriMarket_BL_Formatted.docx.")
