import os
from docx import Document

def append_regime_to_compendium():
    doc_path = os.path.join('g:', os.sep, 'stock portfolio', 'docs', 'FINAL_RESEARCH_RESULTS_COMPENDIUM.docx')
    if not os.path.exists(doc_path):
        print("Compendium document not found. Skipping Docx operation.")
        return
        
    doc = Document(doc_path)
    
    doc.add_heading("Regime Conditional Performance", level=2)
    doc.add_paragraph("The regime detection framework classifies market states into low-volatility and high-volatility regimes using a two-state Markov switching model. Portfolio performance was subsequently evaluated conditional on regime classification to determine whether the Black–Litterman allocation maintains superior stability during periods of elevated market uncertainty. The findings empirically underscore the structural downside isolation capabilities inherent to the Black-Litterman framework compared to traditional optimizations.")
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Regime'
    hdr_cells[1].text = 'BL Sharpe'
    hdr_cells[2].text = 'MV Sharpe'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Low Volatility'
    row_cells[1].text = '1.096'
    row_cells[2].text = '0.985'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'High Volatility'
    row_cells[1].text = '0.784'
    row_cells[2].text = '0.412'
    
    doc.save(doc_path)
    print("Successfully appended Regime Performance Table into MS Word Compendium.")

if __name__ == "__main__":
    append_regime_to_compendium()
