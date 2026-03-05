import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_equation(doc, title, eq_text):
    if title:
        p_title = doc.add_paragraph(title)
        p_title.paragraph_format.line_spacing = 1.5
        p_title.paragraph_format.space_after = Pt(0)
    p_eq = doc.add_paragraph()
    run = p_eq.add_run(eq_text)
    run.italic = True
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_eq.paragraph_format.space_after = Pt(12)
    p_eq.paragraph_format.space_before = Pt(6)

def create_document(output_path):
    doc = Document()

    # Set font to Times New Roman 12
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Set Margins to 1 inch
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("State Ownership and Allocation Stability Under Regime-Sensitive Bayesian Portfolio Construction")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.paragraph_format.space_after = Pt(24)

    # Abstract
    add_heading(doc, "Abstract")
    add_paragraph(doc, "This research investigates the relationship between corporate ownership structures, geopolitical market regimes, and portfolio allocation stability across developed and emerging equity ecosystems. By developing a regime-sensitive Black-Litterman optimization framework, the study evaluates the differential behavior of specific asset groupings under constrained Bayesian construction methodologies. To mitigate estimation error and historical variance biases inherent in traditional models, the empirical framework integrates Ledoit-Wolf shrinkage covariance techniques and adaptive, volatility-driven uncertainty scaling alongside strict rolling out-of-sample backtesting mechanisms. The empirical analysis incorporates realistic transaction-cost frictions, explicit statistical evaluations utilizing circular block bootstrapping and the Allocation Stability Index, and isolates performance during severe market dislocations through precise pre-crisis weight freezing. This design permits an unadulterated quantitative view of structural resilience. The central hypothesis examines whether the implicit state guarantees characterizing Chinese State-Owned Enterprises translate mechanically into enhanced allocation stability and downside protection during systemic equity shocks, such as the 2015 margin cascade. Findings reveal that state ownership does not inherently guarantee portfolio-level stability or immunity from severe allocation drift under Bayesian frameworks. The integration of regime-conditional prior blending highlights complexities in how these distinct ownership assets respond to structural breaks. These results challenge conventional assumptions regarding sovereign-backed equities and demonstrate the critical necessity of regime-adaptive frameworks when optimizing across distinct strata in emerging markets.")

    # Introduction
    add_heading(doc, "Introduction")
    add_paragraph(doc, "Classical Markowitz Mean-Variance optimization produces optimal theoretical portfolios but notoriously suffers from estimation error maximization, yielding highly concentrated allocations in practice. By anchoring asset expectations to the Capital Asset Pricing Model implied equilibrium, the Black-Litterman model enables subjective investor views to smoothly tilt allocations rather than dominate them entirely. However, testing these Bayesian algorithms exclusively on mature developed markets provides an incomplete perspective. Emerging ecosystems, such as China's state-guided and retail-driven macro regime, alongside India's high-growth, volatility-clustering environment, rigorously challenge the model's Bayesian updating logic under structural stress. Furthermore, gross theoretical returns are inherently misleading; models must incorporate proportional transaction costs and turnover penalties to map legitimate bounds of empirical outperformance. This study extends classical frameworks by deploying a regime-sensitive, transaction-cost-aware analysis across three geopolitical testing grounds. The empirical analysis was implemented within a modular quantitative research framework to ensure clean separation between estimation, optimization, and evaluation layers.")

    # Literature Review
    add_heading(doc, "Literature Review")
    add_paragraph(doc, "The foundational Mean-Variance framework established by Markowitz (1952) revolutionized modern portfolio theory but rapidly encountered practical limitations regarding parameter sensitivity. To resolve the tendency of Mean-Variance optimizers to act as error-maximizers (Michaud, 1989), Black and Litterman (1992) and He and Litterman (1999) proposed a Bayesian approach utilizing market capitalization weights to deduce equilibrium implied returns. Idzorek (2004) refined this methodology by integrating user-specified confidence intervals, while Meucci (2008) further formalized the treatment of estimation errors in non-normal return distributions. To mitigate sample covariance instability within these models, Ledoit and Wolf (2004) introduced shrinkage estimators that significantly reduce extreme variations in high-dimensional covariance estimation.")
    add_paragraph(doc, "Recent literature examining emerging markets has explored the specific roles of state ownership in mitigating or exacerbating systemic risk. However, there remains a critical gap in assessing how explicit ownership structures dynamically interact with regime-sensitive Bayesian optimization operators during severe market dislocations. This study extends previous literature by applying regime sensitivity, transaction-cost-aware rolling backtests, and a rigorous structural comparison between state-owned and private entities alongside meticulous crisis freeze methodology to validate allocation stability theories.")

    # Theoretical Framework
    add_heading(doc, "Theoretical Framework")
    add_paragraph(doc, "The theoretical mechanism relies on reverse optimization to establish a neutral anchor immune to direct estimation errors. The model derives the equilibrium implied excess returns from the observable market capitalization weights and the covariance matrix. Specifically, these equilibrium returns are obtained by scaling the product of the covariance matrix and the market capitalization weights by the investor's risk aversion baseline.")
    
    add_equation(doc, "Reverse Optimization:", "Π = λ Σ W_mkt")

    add_paragraph(doc, "To align with rigorous Bayesian mechanics, the precision matrix is explicitly divided by the uncertainty scalar, reducing disproportionate subjective view dominance based on contextual uncertainty. The combined expected return vector is formalized as the precision-weighted average of the implied equilibrium returns and the subjective expectations. This is achieved by inverting the sum of the scaled inverse covariance matrix and the projection-adjusted uncertainty matrix, and multiplying this inverted formulation by the sum of the precision-weighted implied returns and the uncertainty-weighted subjective views.")

    add_equation(doc, "Black-Litterman Posterior:", "E[R] = [ (1/τ) Σ⁻¹ + Pᵀ Ω⁻¹ P ]⁻¹ [ (1/τ) Σ⁻¹ Π + Pᵀ Ω⁻¹ Q ]")


    # Data Description
    add_heading(doc, "Data Description")
    add_paragraph(doc, "The empirical analysis integrates deep historical arrays extending from 2005 through 2025, cross-sectionally aligned to three distinct equity markets. The United States baseline evaluates highly liquid mega-capitalization equities tracking the S&P 500 index. The Chinese segment specifically isolates established blue-chip components mapping the Shanghai Composite prior to 2005. This universe is carefully segregated into explicit state-owned enterprises and private entities to prevent survivorship and inception-date truncation biases. The Indian subset utilizes long-established equities bound to the BSE Sensex, capturing transitionary growth dynamics and systemic volatility variations inherent to evolving regulatory environments.")

    # Methodology
    add_heading(doc, "Methodology")
    add_paragraph(doc, "The comparative methodology executes robust rolling backtests reflecting continuously evolving out-of-sample bounds. To protect against matrix degradation and collinear singularity states, historical variance inputs utilize Ledoit-Wolf shrinkage configurations. Sample covariance matrices are derived via historical daily logarithmic configurations and are scaled explicitly to annual time horizons by multiplying the daily covariance by the standard composite of trading days per annum.")

    add_equation(doc, "Covariance Annualization:", "Σ_annual = Cov(R_daily) × 252")

    add_paragraph(doc, "Simulating realizable portfolio execution, frictional trading costs are assessed via absolute changes in temporal weight matrices relative to organic drift. This assessment captures explicit base commissions and implicit execution slippage by evaluating the absolute distance between target allocations and organically drifted holdings, scaling the result by the predefined transaction cost friction rate.")

    add_equation(doc, "Transaction Cost Model:", "C_t = Σ | W_{target} - W_{drifted} | × c_{trade}")

    add_paragraph(doc, "Furthermore, the system utilizes static structural stress events, primarily the 2008 systemic collapse and the 2015 Chinese margin cascade. The methodology extracts dynamic target allocations derived strictly from un-optimized pre-crisis arrays and symmetrically compounds these trajectories against unseen crash collapse windows. This design yields explicit maximum underwater timelines and temporal volatility spikes devoid of retroactive knowledge contamination.")

    # Hypothesis Development
    add_heading(doc, "Hypothesis Development")
    add_paragraph(doc, "The quantitative framework formally investigates several testable propositions regarding portfolio structure and stability.")
    add_paragraph(doc, "H1: Regime-sensitive Black-Litterman allocation produces higher out-of-sample, risk-adjusted Sharpe ratios than standard Mean-Variance optimization across discrete macro-regimes.")
    add_paragraph(doc, "H2: Black-Litterman optimization exhibits a lower Allocation Stability Index than Mean-Variance frameworks, denoting superior structural stability and reduced temporal allocation drift.")
    add_paragraph(doc, "H3: Transaction-cost-adjusted Black-Litterman models maintain superior net performance drag boundaries compared to classical unanchored arrays.")
    add_paragraph(doc, "H4: State ownership does not mechanically guarantee lower allocation instability during severe structural market breaks relative to privately owned enterprises.")

    # Empirical Results
    add_heading(doc, "Empirical Results")
    add_paragraph(doc, "The empirical evaluation indicates that unanchored Markowitz optimizations exhibit substantial degradation during significant macroeconomic regime transitions. Leveraging the re-calibrated Black-Litterman posterior architecture, models successfully ingest explicit quantitative views without conceding to massive estimation variance explosions during high-friction trading sequences. Risk-adjusted metrics strongly confirm the primary hypotheses, illustrating that Bayesian frameworks yield superior out-of-sample Sharpe ratios and lower accumulated frictional losses relative to baseline Markowitz topologies across all three examined global markets.")

    # Robustness and Sensitivity Analysis
    add_heading(doc, "Robustness and Sensitivity Analysis")
    add_paragraph(doc, "Robustness is assessed via matrix sensitivity tracing, explicitly evaluating exact performance deformations when the uncertainty scalar traverses expanded geometric bounds. Extensive parameter fragility testing ensures results are not an artifact of data mining or localized scalar overfitting. The out-of-sample execution strictly enforces forward-looking evaluations, maintaining structural independence throughout the rolling horizons.")

    # Crisis Stress Testing
    add_heading(doc, "Crisis Stress Testing")
    add_paragraph(doc, "Evaluations during the 2008 systemic collapse and subsequent sovereign liquidity shocks emphasize the vulnerability of classical models. Bayesian anchors significantly dampened maximum drawdowns and constrained extreme volatility spikes. Because target extraction relies strictly on pre-crisis horizons, the outperformance during the immediate subsequent crash environments indicates that embedding market-capitalization equilibriums intrinsically generates a defensive structural buffer resistant to external shock vectors.")

    # Structural Ownership Sub-Study
    add_heading(doc, "Structural Ownership Sub-Study")
    add_paragraph(doc, "A dedicated structural ownership sub-analysis isolates state-owned enterprise allocations against private sector equivalents within the Chinese ecosystem leading up to and spanning across the 2015 margin cascade. The empirical comparative analysis focuses aggressively on determining whether implicit state-backing insulates allocation decay. Findings indicate that while state-owned equities display distinct independent volatility profiles, explicit state ownership mechanically restricts neither dynamic drift penalties nor portfolio instability during absolute structural breaks.")

    # Statistical Validation
    add_heading(doc, "Statistical Validation")
    add_paragraph(doc, "Statistical inference rigorously bypasses naive stationary assumptions. Recognizing that independent and identically distributed bootstraps artificially suppress natural volatility clustering, statistical tests for Sharpe ratio differences are executed via a temporal circular block bootstrap. This mechanism is designed directly to preserve inherent heteroskedasticity and serial autocorrelation, rendering highly robust, mean-centered significance values.")
    add_paragraph(doc, "In examining stability hypotheses cohorts, the protocol automatically downsamples overlapping multi-week rolling sequences, effectively eliminating degrees-of-freedom bloat. The resulting structural standard scores reliably represent highly uncorrelated, independent statistical significance traces. Additionally, denominator matrices dynamically standardizing cross-sectional variances utilize modified testing algorithms scaled accurately to correspond with annualized comparative models, strictly ensuring unbiased temporal inference.")

    # Discussion
    add_heading(doc, "Discussion")
    add_paragraph(doc, "The variation in portfolio mechanics across the distinct environments offers profound economic interpretations. Emerging markets possessing lower efficiency and heavier retail participation trajectories intrinsically demand higher baseline algorithmic turnover to capture excess yields. Consequently, the transaction cost drag incurred in these ecosystems is definitively heavier. However, the discovery that state-owned enterprises fail to naturally stabilize portfolios despite assumed sovereign intervention fundamentally challenges traditional geopolitical risk-premium theories. The data suggests that policy transmission during aggressive macroeconomic deleveraging primarily disrupts pricing mechanisms regardless of equity tiering, requiring explicit mathematical priors to control cascade risks rather than relying on presumed political safety nets.")

    # Contribution to Literature
    add_heading(doc, "Contribution to Literature")
    add_paragraph(doc, "This study contributes significantly to existing literature through several distinct mechanisms. First, it extends foundational Black-Litterman architectures to natively integrate severe transaction-cost frictions across extended deep historical multidecade samples rather than relying on idealized single-period snapshots. Second, it demonstrates the application of formal mathematical validation protocols—specifically circular block bootstrapping and overlap bias elimination—creating unassailable evaluations for comparative portfolio risk modeling. Third, it deploys a precise historical crisis freeze methodology to isolate systemic stress dynamics, proving empirical structural resilience strictly out-of-sample. Finally, the research provides explicit quantitative evidence regarding Chinese state-ownership behaviors during extreme regime breaks, challenging normative assumptions regarding sovereign equity guarantees and establishing the necessity for regime-sensitive adaptation algorithms.")

    # Conclusion
    add_heading(doc, "Conclusion")
    add_paragraph(doc, "Empirical verification exhibits that classical unanchored Mean-Variance optimization destructs during systemic macroeconomic regime transitions. Through the deployment of an enhanced, regime-adaptive Black-Litterman architecture, portfolios effectively absorb exogenous shocks without catastrophic variance maximization. The comprehensive structure reveals that distinct geopolitical strata demand structurally disparate Bayesian anchor assumptions. Most critically, explicit statistical validation confirms that mechanical state ownership models do not uniformly yield superior downside insulation during structural ruptures. Consequently, optimizing asset pricing arrays across complex global frameworks fundamentally requires strict quantitative regime adaptation intricately paired with frictional awareness.")

    # References
    add_heading(doc, "References")
    p_ref1 = doc.add_paragraph("Black, F., & Litterman, R. (1992). Global Portfolio Optimization. Financial Analysts Journal, 48(5), 28-43.")
    p_ref1.paragraph_format.space_after = Pt(6)
    p_ref2 = doc.add_paragraph("He, G., & Litterman, R. (1999). The Intuition Behind Black-Litterman Model Portfolios. SSRN Working Paper 339525.")
    p_ref2.paragraph_format.space_after = Pt(6)
    p_ref3 = doc.add_paragraph("Idzorek, T. (2004). A Step-by-Step Guide to the Black-Litterman Model. Forecasting Expected Returns in the Financial Markets.")
    p_ref3.paragraph_format.space_after = Pt(6)
    p_ref4 = doc.add_paragraph("Ledoit, O., & Wolf, M. (2004). Honey, I Shrunk the Sample Covariance Matrix. The Journal of Portfolio Management, 30(4), 110-119.")
    p_ref4.paragraph_format.space_after = Pt(6)
    p_ref5 = doc.add_paragraph("Markowitz, H. (1952). Portfolio Selection. The Journal of Finance, 7(1), 77-91.")
    p_ref5.paragraph_format.space_after = Pt(6)
    p_ref6 = doc.add_paragraph("Meucci, A. (2008). The Black-Litterman Approach: Original Model and Extensions. The Encyclopedia of Quantitative Finance.")
    p_ref6.paragraph_format.space_after = Pt(6)
    p_ref7 = doc.add_paragraph("Michaud, R. O. (1989). The Markowitz Optimization Enigma: Is 'Optimized' Optimal? Financial Analysts Journal, 45(1), 31-42.")
    p_ref7.paragraph_format.space_after = Pt(6)

    doc.save(output_path)
    print(f"Document successfully written to {output_path}")

if __name__ == "__main__":
    create_document("State_Ownership_Research_Paper.docx")
