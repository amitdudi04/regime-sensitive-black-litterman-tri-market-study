from docx import Document
from docx.shared import Pt
import re

def set_normal_font(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = None # default black
        run.bold = True

def add_paragraph_with_bolding(doc, text):
    p = doc.add_paragraph()
    
    # Simple markdown bold parser
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)

def run():
    doc = Document()
    set_normal_font(doc)

    title = doc.add_heading("Project Architecture Refactoring Plan", level=1)
    for run_obj in title.runs:
        run_obj.font.name = 'Times New Roman'
        run_obj.font.color.rgb = None
        run_obj.bold = True

    add_heading(doc, "Goal Description", 2)
    doc.add_paragraph("The objective is to restructure the existing monolithic and disorganized Regime-Sensitive Black–Litterman Tri-Market research project into a clean, reproducible, and highly modular quantitative research pipeline suitable for academic review. The mathematical logic, numerical results, and model implementations will not be changed.")

    add_heading(doc, "Proposed Changes", 2)
    doc.add_paragraph("We will reorganize the codebase into a clean directory structure with single-responsibility modules. Old monolithic scripts will be archived, not deleted.")

    add_heading(doc, "Directory Structure & New Modules", 3)
    
    add_heading(doc, "config/", 4)
    add_paragraph_with_bolding(doc, "• [NEW] config/project_config.yaml: Store project parameters like start dates, transaction cost rates, tau parameter, and rebalance windows for full reproducibility.")

    add_heading(doc, "data/", 4)
    doc.add_paragraph("• raw/ - Raw market data\n• processed/ - Cleaned datasets")

    add_heading(doc, "core/", 4)
    add_paragraph_with_bolding(doc, "• [NEW] core/data_loader.py: Download market data via yfinance, handle MultiIndex safely, return clean price DataFrames.")
    add_paragraph_with_bolding(doc, "• [NEW] core/return_calculations.py: Compute log returns, handle missing data/NaNs.")
    add_paragraph_with_bolding(doc, "• [NEW] core/covariance_estimators.py: Implement Ledoit–Wolf shrinkage and annualized covariance calculations.")

    add_heading(doc, "models/", 4)
    add_paragraph_with_bolding(doc, "• [NEW] models/black_litterman_model.py: Compute equilibrium returns and implement the Bayesian posterior update logic.")
    add_paragraph_with_bolding(doc, "• [NEW] models/optimizer.py: Compute optimal weight allocations (Mean-Variance and unconstrained optimizers).")

    add_heading(doc, "backtesting/", 4)
    add_paragraph_with_bolding(doc, "• [NEW] backtesting/rolling_backtest.py: Perform rolling out-of-sample testing logic.")
    add_paragraph_with_bolding(doc, "• [NEW] backtesting/transaction_costs.py: Compute turnover and apply frictional decay penalties to derive net returns.")
    add_paragraph_with_bolding(doc, "• [NEW] backtesting/crisis_freeze.py: Implement the crisis freeze methodology (e.g., locking weights before 2008/2015/2020 events).")
    add_paragraph_with_bolding(doc, "• [NEW] backtesting/allocation_stability_index.py: Compute the ASI metric across rolling rebalance blocks.")

    add_heading(doc, "analysis/", 4)
    add_paragraph_with_bolding(doc, "• [NEW] analysis/soe_private_analysis.py: Split Chinese equities into SOE and Private categorizations and compute independent metrics.")
    add_paragraph_with_bolding(doc, "• [NEW] analysis/statistical_tests.py: Implement Circular Block Bootstrapping, T-tests, and Jobson-Korkie standardizations.")

    add_heading(doc, "pipelines/ (High-Level Execution)", 4)
    add_paragraph_with_bolding(doc, "• [NEW] pipelines/run_tri_market_pipeline.py: Loads US/China/India data, runs Black-Litterman vs Markowitz, rolling backtests, transaction costs, and exports full results.")
    add_paragraph_with_bolding(doc, "• [NEW] pipelines/run_soe_pipeline.py: Loads China data, separates SOE/Private, runs allocations, ASI, stats, and exports comparison tables.")
    add_paragraph_with_bolding(doc, "• [NEW] pipelines/run_crisis_analysis.py: Extracts pre-crisis weights, simulates crashes, computes drawdowns/vol spikes, and exports tables.")

    add_heading(doc, "experiments/", 4)
    add_paragraph_with_bolding(doc, "• [NEW] experiments/run_tau_sensitivity.py: Robustness testing for the tau parameter.")

    add_heading(doc, "results/", 4)
    doc.add_paragraph("• Outputs: tri_market_summary.csv, soe_vs_private_china_ownership_summary.csv, crisis_comparison_tables.csv, statistical_test_output.json")
    add_paragraph_with_bolding(doc, "• [NEW] results/export_utils.py: Utility functions to export clean .csv and .json tables matching research paper outputs.")

    add_heading(doc, "visualization/", 4)
    add_paragraph_with_bolding(doc, "• [NEW] visualization/plotting_tools.py: Isolate plotting capabilities.")

    add_heading(doc, "tests/", 4)
    add_paragraph_with_bolding(doc, "• [NEW] tests/test_data_loader.py")
    add_paragraph_with_bolding(doc, "• [NEW] tests/test_black_litterman.py")
    add_paragraph_with_bolding(doc, "• [NEW] tests/test_backtest.py")
    add_paragraph_with_bolding(doc, "• [NEW] tests/test_optimizer.py")

    add_heading(doc, "legacy/", 4)
    add_paragraph_with_bolding(doc, "• [MOVE] Archive all monolithic scripts here instead of permanently deleting them.")

    add_heading(doc, "docs/ & Project Root", 4)
    add_paragraph_with_bolding(doc, "• [MOVE] Move FINAL_PROJECT_IMPLEMENTATION.md and FINAL_RESEARCH_RESULTS_COMPENDIUM.docx into docs/.")
    add_paragraph_with_bolding(doc, "• [NEW] README.md: Professional project overview explaining the methodology and exact run commands.")
    add_paragraph_with_bolding(doc, "• [NEW] requirements.txt: Package dependency mapping.")

    add_heading(doc, "Verification Plan", 2)
    
    add_heading(doc, "Automated Tests", 3)
    doc.add_paragraph("1. Execute python pipelines/run_tri_market_pipeline.py and verify that the tri_market_summary.csv matches the existing results.")
    doc.add_paragraph("2. Execute python pipelines/run_soe_pipeline.py and verify soe_vs_private_china_ownership_summary.csv accuracy.")
    doc.add_paragraph("3. Execute python pipelines/run_crisis_analysis.py and verify the crisis_comparison_tables.csv.")

    add_heading(doc, "Manual Verification", 3)
    doc.add_paragraph("Review the exported CSV and JSON results against the values defined in docs/FINAL_PROJECT_IMPLEMENTATION.md to guarantee structural mathematical preservation.")

    doc.save("Project_Architecture_Refactoring_Plan.docx")
    print("Implementation plan successfully converted to docx.")

if __name__ == "__main__":
    run()
